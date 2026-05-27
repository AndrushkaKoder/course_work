#!/usr/bin/env python3
"""Сборка пояснительной записки в Word: титул (.pages) + текст + ГОСТ."""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docxcompose.composer import Composer

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "POYASNITELNAYA-ZAPISKA-full.md"
TITLE_PAGES = ROOT / "Титульный лист.pages"
TITLE_DOCX = ROOT / "_title_temp.docx"
BODY_DOCX = ROOT / "_body_temp.docx"
OUTPUT = ROOT / "POYASNITELNAYA-ZAPISKA-full.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
FIRST_LINE_INDENT = Cm(1.25)
HANGING_INDENT = Cm(1.25)
PLACEHOLDER_TEXT = "Место для скриншота / схемы таблицы"

FIGURE_CAPTIONS: dict[str, str] = {
    "1.1": "Контекстная IDEF0-диаграмма «Управление деятельностью автосалона»",
    "1.2": "Декомпозиция функций автосалона (первый уровень)",
    "2.1": "Контекстная диаграмма потоков данных системы DriveLine",
    "2.2": "Диаграмма потоков данных (первый уровень декомпозиции)",
    "2.3": "Концептуальная ER-модель базы данных",
    "2.4": "Согласованные состояния сделки продажи и автомобиля на складе",
    "3.1": "Архитектура развёртывания информационной системы DriveLine",
    "3.2": "Алгоритм синхронизации статуса автомобиля при сделке продажи",
    "3.3": "Алгоритм формирования отчёта в фоновом режиме",
}

FIGURE_REF = re.compile(
    r"(?:рисунок|рисунке|рис\.)\s*([\d.]+)|`diagrams/([^`]+\.mmd)`",
    re.IGNORECASE,
)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size=None) -> None:
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.insert(0, r_fonts)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = FIRST_LINE_INDENT

    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = Pt(14)
        style.font.bold = True
        hpf = style.paragraph_format
        hpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        hpf.space_before = Pt(12 if level == 1 else 6)
        hpf.space_after = Pt(6)
        hpf.first_line_indent = Cm(0)


def add_page_number_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    set_run_font(run)
    for part, ftype in (
        ("begin", "begin"),
        (None, "instr"),
        ("separate", "separate"),
        (None, "text"),
        ("end", "end"),
    ):
        if ftype == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)
        elif ftype == "text":
            t = OxmlElement("w:t")
            t.text = "2"
            run._r.append(t)
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), part)
            run._r.append(el)


def add_toc_field(doc: Document) -> None:
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.first_line_indent = Cm(0)
    run = h.add_run("СОДЕРЖАНИЕ")
    set_run_font(run, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    set_run_font(run)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    r = note.add_run(
        "(После открытия файла в Word: правый клик по оглавлению → «Обновить поле» → «Обновить целиком».)"
    )
    set_run_font(r, italic=True, size=Pt(12))
    doc.add_page_break()


def export_title_from_pages() -> None:
    if not TITLE_PAGES.is_file():
        raise FileNotFoundError(f"Не найден титульный лист: {TITLE_PAGES}")
    script = f'''
    set pagesFile to POSIX file "{TITLE_PAGES}"
    set outFile to POSIX file "{TITLE_DOCX}"
    tell application "Pages"
        set theDoc to open pagesFile
        delay 1
        export theDoc to outFile as Microsoft Word
        close theDoc saving no
    end tell
    '''
    subprocess.run(["osascript", "-e", script], check=True, capture_output=True, text=True)


def ensure_title_docx() -> Path:
    if not TITLE_DOCX.is_file() or (
        TITLE_PAGES.is_file() and TITLE_PAGES.stat().st_mtime > TITLE_DOCX.stat().st_mtime
    ):
        export_title_from_pages()
    return TITLE_DOCX


def add_figure_placeholder(doc: Document, figure_id: str) -> None:
    if figure_id in getattr(doc, "_figures_added", set()):
        return
    if not hasattr(doc, "_figures_added"):
        doc._figures_added = set()
    doc._figures_added.add(figure_id)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.space_before = Pt(6)

    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.first_line_indent = Cm(0)
    run = box.add_run(f"*{PLACEHOLDER_TEXT}*")
    set_run_font(run, italic=True)

    caption = FIGURE_CAPTIONS.get(figure_id, "Иллюстрация")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(f"Рисунок {figure_id} – {caption}")
    set_run_font(run)


def extract_figure_ids(text: str) -> list[str]:
    ids: list[str] = []
    for num, _path in FIGURE_REF.findall(text):
        if num:
            ids.append(num.strip())
    return ids


def add_inline_runs(paragraph, text: str, *, base_bold: bool = False) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, bold=base_bold)
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, bold=True)
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run)
        elif chunk.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if link_match:
                label, url = link_match.groups()
                run = paragraph.add_run(f"{label} ({url})")
                set_run_font(run, bold=base_bold)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, bold=base_bold)


def add_paragraph(doc: Document, text: str, *, indent: bool = True, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    add_inline_runs(p, text.strip())
    if not p.runs:
        set_run_font(p.add_run(text.strip()))
    for fig_id in extract_figure_ids(text):
        add_figure_placeholder(doc, fig_id)
    if re.search(r"скриншот", text, re.I) and re.search(r"приложени|форм moonshine|витрин", text, re.I):
        add_ui_screenshot_placeholder(doc)


def add_ui_screenshot_placeholder(doc: Document) -> None:
    key = "ui-screenshot"
    if key in getattr(doc, "_figures_added", set()):
        return
    doc._figures_added.add(key)
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.first_line_indent = Cm(0)
    run = box.add_run(f"*{PLACEHOLDER_TEXT}*")
    set_run_font(run, italic=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    clean = text.strip()
    if level == 1 and re.match(r"^\d+\s", clean):
        clean = clean.upper()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean)
        set_run_font(run, bold=True)
        return
    p = doc.add_heading(clean, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, bold=True)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    body = rows[2:] if len(rows) > 1 and re.match(r"^[-:|\s]+$", "|".join(rows[1])) else rows[1:]

    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        add_inline_runs(p, header)
        for run in p.runs:
            set_run_font(run, bold=True)

    for row_idx, row in enumerate(body):
        for col, value in enumerate(row):
            if col >= len(headers):
                break
            cell = table.rows[row_idx + 1].cells[col]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(p, value)
            for run in p.runs:
                set_run_font(run)

    doc.add_paragraph().paragraph_format.first_line_indent = Cm(0)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    text = "\n".join(lines)
    run = p.add_run(text)
    set_run_font(run, size=Pt(12))


def add_list_item(doc: Document, text: str, *, ordered: bool, number: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    prefix = f"{number}. " if ordered and number else "– "
    run = p.add_run(prefix)
    set_run_font(run)
    add_inline_runs(p, text)
    for fig_id in extract_figure_ids(text):
        add_figure_placeholder(doc, fig_id)


def add_bibliography_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = HANGING_INDENT
    add_inline_runs(p, text)


def build_body(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    doc._figures_added = set()
    add_toc_field(doc)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_bibliography = False
    i = 0

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            i += 1
            continue

        if line.strip().startswith("|"):
            table_rows.append(parse_table_row(line))
            i += 1
            continue
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if "список использованных источников" in title.lower():
                in_bibliography = True
            add_heading(doc, title, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            num, _, rest = stripped.partition(". ")
            if in_bibliography:
                add_bibliography_item(doc, stripped)
            else:
                add_list_item(doc, rest, ordered=True, number=num)
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:].strip(), ordered=False)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    if table_rows:
        add_table(doc, table_rows)

    doc.save(out_path)


def apply_page_numbering(final_path: Path) -> None:
    """Титул — без номера; со 2-й страницы (оглавление) — сквозная нумерация с «2»."""
    doc = Document(final_path)
    for section in doc.sections:
        sect_pr = section._sectPr
        title_pg = sect_pr.find(qn("w:titlePg"))
        if title_pg is None:
            title_pg = OxmlElement("w:titlePg")
            sect_pr.append(title_pg)
        title_pg.set(qn("w:val"), "1")

        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)
        pg_num_type.set(qn("w:start"), "2")

        section.footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False

        for footer in (section.footer, section.first_page_footer):
            for p in list(footer.paragraphs):
                p._element.getparent().remove(p._element)

        add_page_number_field(section.footer.add_paragraph())

    doc.save(final_path)


def merge_documents(title_path: Path, body_path: Path, output_path: Path) -> None:
    composer = Composer(Document(title_path))
    composer.append(Document(body_path))
    composer.save(output_path)


def build(md_path: Path, output_path: Path) -> None:
    title_path = ensure_title_docx()
    build_body(md_path, BODY_DOCX)
    merge_documents(title_path, BODY_DOCX, output_path)
    apply_page_numbering(output_path)


def main() -> int:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    if not source.is_file():
        print(f"Файл не найден: {source}", file=sys.stderr)
        return 1
    build(source, output)
    print(f"Создан: {output}")
    print("Откройте в Word и обновите оглавление (правый клик → Обновить поле).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
